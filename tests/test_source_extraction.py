from __future__ import annotations

import io
import sys
from types import SimpleNamespace

import pytest
from docx import Document

from app.services import text_extraction
from app.services.text_extraction import (
    MAX_EXTRACTED_CHARS,
    PDF_PAGE_SEPARATOR,
    ExtractionResult,
    extract_source,
    extract_text,
    extract_text_result,
)


def test_plain_text_is_normalized_and_sections_are_detected() -> None:
    content = (
        b"\xef\xbb\xbfSUMMARY\r\n\r\n\r\nBuilt cafe\xcc\x81 tools.\x00\r"
        b"EDUCATION\r\nTexas Tech University\r\n"
    )

    result = extract_text_result("resume.txt", content)

    assert isinstance(result, ExtractionResult)
    assert result.status == "ok"
    assert result.kind == "txt"
    assert result.bytes == len(content)
    assert result.pages is None
    assert result.text == (
        "SUMMARY\n\nBuilt caf\u00e9 tools.\nEDUCATION\nTexas Tech University"
    )
    assert result.extracted_chars == len(result.text)
    assert result.words == 8
    assert result.detected_sections == ["Summary", "Education"]
    assert any("NUL" in warning for warning in result.warnings)


def test_invalid_utf8_uses_visible_controlled_fallback() -> None:
    result = extract_source("notes.md", b"PROJECTS\nBuilt Stripe \xff integration")

    assert result.status == "ok"
    assert "\ufffd" in result.text
    assert any("not valid UTF-8" in warning for warning in result.warnings)
    assert result.detected_sections == ["Projects"]


def test_docx_preserves_paragraph_and_table_order() -> None:
    document = Document()
    document.add_paragraph("EXPERIENCE")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Company"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Acme"
    table.cell(1, 1).text = "Engineer"
    document.add_paragraph("EDUCATION")
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract_text_result("resume.DOCX", buffer.getvalue())

    assert result.status == "ok"
    assert result.pages is None
    assert result.text == (
        "EXPERIENCE\n\nCompany\tRole\nAcme\tEngineer\n\nEDUCATION"
    )
    assert result.detected_sections == ["Experience", "Education"]


def test_pdf_keeps_boundaries_and_warns_about_sparse_pages(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pages = [
        SimpleNamespace(extract_text=lambda: "EXPERIENCE\n" + "A" * 50),
        SimpleNamespace(extract_text=lambda: ""),
        SimpleNamespace(extract_text=lambda: "EDUCATION\n" + "B" * 50),
    ]
    fake_module = SimpleNamespace(
        PdfReader=lambda _stream: SimpleNamespace(pages=pages)
    )
    monkeypatch.setitem(sys.modules, "pypdf", fake_module)

    result = extract_text_result("resume.pdf", b"%PDF-fake")

    assert result.status == "ok"
    assert result.pages == 3
    assert result.text.count(PDF_PAGE_SEPARATOR) == 2
    assert result.text.split(PDF_PAGE_SEPARATOR)[1] == ""
    assert result.detected_sections == ["Experience", "Education"]
    assert any(
        "page(s) 2" in warning and "OCR" in warning
        for warning in result.warnings
    )


def test_empty_pdf_fails_with_page_and_ocr_metadata(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pages = [SimpleNamespace(extract_text=lambda: None)]
    fake_module = SimpleNamespace(
        PdfReader=lambda _stream: SimpleNamespace(pages=pages)
    )
    monkeypatch.setitem(sys.modules, "pypdf", fake_module)

    result = extract_text_result("scan.pdf", b"%PDF-fake")

    assert result.status == "failed"
    assert result.pages == 1
    assert result.text == ""
    assert any("OCR" in warning for warning in result.warnings)


def test_truncation_metadata_matches_backward_compatible_output(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(text_extraction, "MAX_EXTRACTED_CHARS", 25)
    content = b"SUMMARY\n" + (b"word " * 20)

    result = extract_text_result("long.txt", content)
    text, status = extract_text("long.txt", content)

    assert result.status == status == "ok"
    assert result.truncated is True
    assert result.extracted_chars == len(content.decode("utf-8").strip())
    assert len(result.text) <= 25
    assert text == result.text
    assert any("truncated" in warning for warning in result.warnings)


def test_legacy_api_and_unsupported_status_remain_compatible() -> None:
    text, status = extract_text("resume.rtf", b"{rtf}")
    result = extract_text_result("resume.rtf", b"{rtf}")

    assert (text, status) == ("", "unsupported")
    assert result.kind == "unsupported"
    assert result.status == "unsupported"
    assert result.extracted_chars == 0
    assert result.words == 0
    assert result.truncated is False
